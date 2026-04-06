from flask import Blueprint, request, jsonify
import os
import uuid
import re
import pymysql
from werkzeug.utils import secure_filename
import docx
from docx.shared import Inches
import base64
from io import BytesIO

contracts_bp = Blueprint('contracts', __name__)

# 上传目录
UPLOAD_FOLDER = 'uploads/contracts'
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

# 允许的文件类型
ALLOWED_EXTENSIONS = {'doc', 'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def format_size(size_in_bytes):
    """格式化文件大小"""
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size_in_bytes < 1024.0:
            return f"{size_in_bytes:.1f}{unit}"
        size_in_bytes /= 1024.0
    return f"{size_in_bytes:.1f}TB"

def extract_docx_content(filepath):
    """提取Word文档内容并保留格式"""
    try:
        doc = docx.Document(filepath)
        html_content = []
        
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                html_content.append('<p>&nbsp;</p>')
                continue
            
            p_html = ['<p>']
            for run in paragraph.runs:
                run_html = run.text
                # 保留格式
                if run.bold:
                    run_html = f'<strong>{run_html}</strong>'
                if run.italic:
                    run_html = f'<em>{run_html}</em>'
                if run.underline:
                    run_html = f'<u>{run_html}</u>'
                p_html.append(run_html)
            p_html.append('</p>')
            html_content.append(''.join(p_html))
        
        return '\n'.join(html_content)
    except Exception as e:
        print(f"提取Word内容失败: {e}")
        return ''

@contracts_bp.route('/api/contracts', methods=['GET'])
def get_contracts():
    """获取合同列表"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor(pymysql.cursors.DictCursor)
        cursor.execute('SELECT id, name, filename, size, status, created_at FROM contracts')
        contracts = cursor.fetchall()
        
        # 格式化创建时间
        for contract in contracts:
            contract['created_at'] = contract['created_at'].strftime('%Y-%m-%d %H:%M:%S')
        
        cursor.close()
        conn.close()
        return jsonify({'success': True, 'data': contracts})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@contracts_bp.route('/api/contracts/upload', methods=['POST'])
def upload_contract():
    """上传合同"""
    try:
        if 'contract' not in request.files:
            return jsonify({'success': False, 'message': '没有文件被上传'}), 400
        
        file = request.files['contract']
        if file.filename == '':
            return jsonify({'success': False, 'message': '请选择文件'}), 400
        
        if not allowed_file(file.filename):
            return jsonify({'success': False, 'message': '只允许上传 .doc 或 .docx 文件'}), 400
        
        # 生成唯一文件名
        original_filename = file.filename
        filename = secure_filename(original_filename)
        unique_filename = f"{uuid.uuid4()}_{filename}"
        filepath = os.path.join(UPLOAD_FOLDER, unique_filename)
        
        # 保存文件
        file.save(filepath)
        
        # 获取文件大小
        file_size = os.path.getsize(filepath)
        formatted_size = format_size(file_size)
        
        # 提取合同名称（去除扩展名）
        base_name = os.path.basename(original_filename)
        contract_name = base_name.rsplit('.', 1)[0] if '.' in base_name else base_name
        
        # 提取Word内容
        content = extract_docx_content(filepath)
        
        # 创建合同记录
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO contracts (name, filename, filepath, size, content, status)
            VALUES (%s, %s, %s, %s, %s, %s)
        ''', (contract_name, filename, filepath, formatted_size, content, 'active'))
        conn.commit()
        cursor.close()
        conn.close()
        
        return jsonify({'success': True, 'message': '合同上传成功'})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@contracts_bp.route('/api/contracts/<int:contract_id>', methods=['GET'])
def get_contract(contract_id):
    """获取合同详情"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor(pymysql.cursors.DictCursor)
        cursor.execute('''
            SELECT id, name, filename, size, content, status, created_at
            FROM contracts
            WHERE id = %s
        ''', (contract_id,))
        contract = cursor.fetchone()
        
        if not contract:
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '合同不存在'}), 404
        
        # 格式化创建时间
        contract['created_at'] = contract['created_at'].strftime('%Y-%m-%d %H:%M:%S')
        
        cursor.close()
        conn.close()
        return jsonify({'success': True, 'data': contract})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@contracts_bp.route('/api/contracts/<int:contract_id>', methods=['PUT'])
def update_contract(contract_id):
    """更新合同内容"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # 检查合同是否存在
        cursor.execute('SELECT id FROM contracts WHERE id = %s', (contract_id,))
        if not cursor.fetchone():
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '合同不存在'}), 404
        
        data = request.get_json()
        if 'content' in data:
            cursor.execute('UPDATE contracts SET content = %s WHERE id = %s', (data['content'], contract_id))
            conn.commit()
        
        cursor.close()
        conn.close()
        return jsonify({'success': True, 'message': '合同更新成功'})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@contracts_bp.route('/api/contracts/<int:contract_id>/status', methods=['PUT'])
def update_contract_status(contract_id):
    """更新合同状态"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        
        # 检查合同是否存在
        cursor.execute('SELECT id FROM contracts WHERE id = %s', (contract_id,))
        if not cursor.fetchone():
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '合同不存在'}), 404
        
        data = request.get_json()
        if 'status' in data:
            status = data['status']
            if status not in ['active', 'inactive']:
                cursor.close()
                conn.close()
                return jsonify({'success': False, 'message': '无效的状态值'}), 400
            
            cursor.execute('UPDATE contracts SET status = %s WHERE id = %s', (status, contract_id))
            conn.commit()
        
        cursor.close()
        conn.close()
        return jsonify({'success': True, 'message': '状态更新成功'})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@contracts_bp.route('/api/contracts/<int:contract_id>', methods=['DELETE'])
def delete_contract(contract_id):
    """删除合同"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor(pymysql.cursors.DictCursor)
        
        # 获取合同信息
        cursor.execute('SELECT filepath FROM contracts WHERE id = %s', (contract_id,))
        contract = cursor.fetchone()
        
        if not contract:
            cursor.close()
            conn.close()
            return jsonify({'success': False, 'message': '合同不存在'}), 404
        
        # 删除文件
        if os.path.exists(contract['filepath']):
            os.remove(contract['filepath'])
        
        # 删除数据库记录
        cursor.execute('DELETE FROM contracts WHERE id = %s', (contract_id,))
        conn.commit()
        
        cursor.close()
        conn.close()
        return jsonify({'success': True, 'message': '合同删除成功'})
    except Exception as e:
        return jsonify({'success': False, 'message': str(e)}), 500

@contracts_bp.route('/api/contracts/sign', methods=['POST'])
def sign_contract():
    """签署合同"""
    try:
        print("收到合同签署请求")
        data = request.get_json()
        if not data:
            print("请求数据为空")
            return jsonify({'success': False, 'message': '请求数据为空'}), 400

        contract_id = data.get('contract_id')
        flow_id = data.get('flow_id')
        id_card = data.get('id_card')
        address = data.get('address')
        contact = data.get('contact')
        signature = data.get('signature')

        print(f"合同ID: {contract_id}, 流程ID: {flow_id}")
        
        if not all([contract_id, flow_id, id_card, address, contact, signature]):
            print("缺少必要参数")
            return jsonify({'success': False, 'message': '缺少必要参数'}), 400

        party_b_name = ''
        try:
            conn_for_flow = get_db_connection()
            cursor_flow = conn_for_flow.cursor(pymysql.cursors.DictCursor)
            cursor_flow.execute("SELECT employee_name FROM flows WHERE flow_id = %s", (flow_id,))
            flow_row = cursor_flow.fetchone()
            if flow_row and flow_row.get('employee_name'):
                party_b_name = flow_row['employee_name']
                print(f"从流程获取到乙方姓名: {party_b_name}")
            cursor_flow.close()
            conn_for_flow.close()
        except Exception as e:
            print(f"获取流程乙方姓名失败（非致命）: {e}")

        # 获取合同模板信息
        try:
            conn = get_db_connection()
            cursor = conn.cursor(pymysql.cursors.DictCursor)
            print("数据库连接成功")
            
            cursor.execute('''
                SELECT name, content, filepath
                FROM contracts
                WHERE id = %s
            ''', (contract_id,))
            contract = cursor.fetchone()
            if not contract:
                print("合同模板不存在")
                cursor.close()
                conn.close()
                return jsonify({'success': False, 'message': '合同模板不存在'}), 404
            print(f"找到合同模板: {contract['name']}")

            # 生成签署后的合同内容（使用占位符覆盖替换）
            signed_content = apply_contract_replacements(
                content=contract['content'],
                party_b_name=party_b_name,
                id_card=id_card,
                phone=contact,
                address=address,
                signature_data=signature
            )

            # 处理签名图片
            try:
                signature_data = signature.split(',')[1]  # 移除base64前缀
                signature_bytes = base64.b64decode(signature_data)
                signature_image = BytesIO(signature_bytes)
                print("签名图片处理成功")
            except Exception as e:
                print(f"签名图片处理失败: {e}")
                # 继续执行，不影响签署记录的保存

            # 生成新的Word文件
            signed_filepath = None
            try:
                signed_filepath = generate_signed_contract(contract['filepath'], id_card, address, contact, signature_image)
                print(f"生成签署合同文件成功: {signed_filepath}")
            except Exception as e:
                print(f"生成签署合同文件失败: {e}")

            # 保存签署记录
            try:
                # 直接使用不包含signed_filepath字段的插入语句
                cursor.execute('''
                    INSERT INTO contract_signatures (contract_id, flow_id, id_card, address, contact, signature, status)
                    VALUES (%s, %s, %s, %s, %s, %s, %s)
                ''', (contract_id, flow_id, id_card, address, contact, signature, 'signed'))
                print("插入签署记录成功")
            except Exception as e:
                print(f"插入签署记录失败: {e}")
                return jsonify({'success': False, 'message': f'保存签署记录失败: {str(e)}'}), 500
            conn.commit()
            cursor.close()
            conn.close()
            print("数据库操作完成")

            return jsonify({'success': True, 'message': '合同签署成功'})
        except Exception as e:
            print(f"数据库操作失败: {e}")
            return jsonify({'success': False, 'message': f'数据库操作失败: {str(e)}'}), 500
    except Exception as e:
        print(f"签署合同失败: {e}")
        return jsonify({'success': False, 'message': f'签署合同失败: {str(e)}'}), 500

@contracts_bp.route('/api/contracts/signed/<int:flow_id>', methods=['GET'])
def get_signed_contract(flow_id):
    """根据流程ID获取已签署的合同"""
    try:
        conn = get_db_connection()
        cursor = conn.cursor(pymysql.cursors.DictCursor)
        
        # 使用流程ID查找合同签署记录
        cursor.execute("""
            SELECT cs.*, c.name as contract_name, c.content, c.filepath as template_file_path
            FROM contract_signatures cs
            LEFT JOIN contracts c ON cs.contract_id = c.id
            WHERE cs.flow_id = %s
            AND cs.status = 'signed'
            ORDER BY cs.signed_at DESC
            LIMIT 1
        """, (flow_id,))
        
        contract = cursor.fetchone()
        
        cursor.close()
        conn.close()
        
        if not contract:
            return jsonify({
                'success': False,
                'message': '未找到已签署的合同'
            }), 404
        
        # 格式化日期
        if contract.get('signed_at'):
            contract['signed_at'] = contract['signed_at'].strftime('%Y-%m-%d %H:%M:%S')
        
        party_b_name = ''
        try:
            conn_flow = get_db_connection()
            cursor_flow = conn_flow.cursor(pymysql.cursors.DictCursor)
            cursor_flow.execute("SELECT employee_name FROM flows WHERE flow_id = %s", (flow_id,))
            flow_row = cursor_flow.fetchone()
            if flow_row and flow_row.get('employee_name'):
                party_b_name = flow_row['employee_name']
            cursor_flow.close()
            conn_flow.close()
        except Exception:
            pass
        
        # 生成已签署合同的HTML内容（使用占位符覆盖替换）
        signed_content = apply_contract_replacements(
            content=contract['content'],
            party_b_name=party_b_name,
            id_card=contract.get('id_card', ''),
            phone=contract.get('contact', ''),
            address=contract.get('address', ''),
            signature_data=contract.get('signature')
        )
        
        # 添加合同头部信息
        final_content = f"""
        <h2>已签署合同</h2>
        <p><strong>合同名称:</strong> {contract.get('contract_name', '未知')}</p>
        <p><strong>签署时间:</strong> {contract.get('signed_at', '未知')}</p>
        <div style="border: 1px solid #ddd; padding: 15px; border-radius: 4px; background-color: #f9f9f9;">
            {signed_content}
        </div>
        """
        
        # 更新content字段为已签署的合同内容
        contract['content'] = final_content
        
        return jsonify({
            'success': True,
            'data': contract
        })
    except Exception as e:
        print(f"获取已签署合同失败: {e}")
        return jsonify({
            'success': False,
            'message': str(e)
        }), 500

# 导入数据库连接函数
def generate_signed_contract(template_path, id_card, address, contact, signature_image):
    """生成签署后的合同Word文件"""
    try:
        # 检查模板文件是否存在
        if not os.path.exists(template_path):
            print(f"模板文件不存在: {template_path}")
            return None
        
        # 打开模板文件
        doc = docx.Document(template_path)
        
        # 遍历所有段落，查找需要替换的内容
        for paragraph in doc.paragraphs:
            # 替换乙方信息
            if '身份证号码 :' in paragraph.text:
                paragraph.text = paragraph.text.replace('身份证号码 :', f'身份证号码 : {id_card}')
            if '送达地址 :' in paragraph.text:
                paragraph.text = paragraph.text.replace('送达地址 :', f'送达地址 : {address}')
            if '电子邮箱/QQ/微信号码 :' in paragraph.text:
                paragraph.text = paragraph.text.replace('电子邮箱/QQ/微信号码 :', f'电子邮箱/QQ/微信号码 : {contact}')
            
            # 在乙方签字位置后添加签名图片
            if '乙方(签字):' in paragraph.text or '乙方 :' in paragraph.text or '乙方:' in paragraph.text or '乙方（签字）' in paragraph.text:
                try:
                    run = paragraph.add_run()
                    run.add_picture(signature_image, width=Inches(2.0))
                except Exception as e:
                    print(f"添加签名图片失败: {e}")
        
        # 生成唯一文件名
        signed_folder = 'uploads/signed_contracts'
        if not os.path.exists(signed_folder):
            try:
                os.makedirs(signed_folder)
            except Exception as e:
                print(f"创建签署合同文件夹失败: {e}")
                return None
        
        unique_filename = f"signed_{uuid.uuid4()}.docx"
        signed_filepath = os.path.join(signed_folder, unique_filename)
        
        # 保存文件
        try:
            doc.save(signed_filepath)
        except Exception as e:
            print(f"保存签署合同失败: {e}")
            return None
        
        return signed_filepath
    except Exception as e:
        print(f"生成签署合同失败: {e}")
        return None

def get_db_connection():
    from config.database import get_db_connection as db_conn
    return db_conn()

def apply_contract_replacements(content, party_b_name='', id_card='', phone='', address='',
                                emergency_name='', emergency_phone='', emergency_address='',
                                signature_data=None):
    if not content:
        return content

    result = content

    for key, value in {
        '__PARTY_B__': f'乙方（承揽人）：{party_b_name}' if party_b_name else '',
        '__ID_CARD__': id_card or '',
        '__PHONE__': phone or '',
        '__ADDRESS__': address or '',
        '__EMERGENCY_NAME__': emergency_name or '',
        '__EMERGENCY_PHONE__': emergency_phone or '',
        '__EMERGENCY_ADDRESS__': emergency_address or '',
    }.items():
        result = result.replace(key, value)
    result = re.sub(r'__(PARTY_B|ID_CARD|PHONE|ADDRESS|EMERGENCY_NAME|EMERGENCY_PHONE|EMERGENCY_ADDRESS)__', '', result)

    if party_b_name:
        result = re.sub(
            r'(乙方)\s+[（(]?(?:承揽人)?[）)]?\s*[：:]\s{2,}(?!.*签字)',
            rf'\1（承揽人）：{party_b_name}', result
        )
        result = re.sub(
            r'(乙方\s*[（(]?)(?:承揽人)?([）)]?\s*[：:])\s*</p>',
            rf'\1\2（承揽人）：{party_b_name}</p>', result
        )

    if id_card:
        def repl_id(m):
            return m.group(1) + id_card
        result = re.sub(
            r'(身份证号(?:码)?[（(]?(?:承揽人)?[）)]?\s*[：:])(?:\s*(?:<u>[^<]*</u>)|[ \t]*\n)',
            repl_id, result
        )
        result = re.sub(
            r'(身份证号码?)\s*[：:][ \t]*$',
            rf'\1：{id_card}', result, flags=re.MULTILINE
        )

    if phone:
        def repl_ph(m):
            return m.group(1) + phone
        result = re.sub(
            r'((?:电话|手机)[（(]?(?:手机)?[）)]?\s*[：:])\s*(<u>[^<]*</u>)',
            repl_ph, result
        )
        result = re.sub(
            r'((?:电话|手机)[（(]?(?:手机)?[）)]?\s*[：:])[ \t]{2,}',
            repl_ph, result
        )

    if address:
        def repl_ad(m):
            return m.group(1) + address
        result = re.sub(
            r'((?:送达)?地址)\s*[：:]\s*(<u>[^<]*</u>)',
            repl_ad, result
        )
        result = re.sub(
            r'((?:送达)?地址)\s*[：:](?=[，。\n]|</p>)',
            repl_ad, result
        )
        result = re.sub(
            r'(送达地址)\s*[：:]',
            repl_ad, result
        )

    contact = phone or address or ''
    if contact:
        result = result.replace('电子邮箱/QQ/微信号码 :', f'电子邮箱/QQ/微信号码：{contact}')
        result = result.replace('电子邮箱/QQ/微信号码:', f'电子邮箱/QQ/微信号码：{contact}')

    if emergency_name:
        result = re.sub(
            r'(紧急联系?人[（(]?\s*姓名[）)]?\s*[：:])\s*(<u>[^<]*</u>|\s+)',
            rf'\1{emergency_name}', result
        )
        result = re.sub(
            r'(姓名)\s*[：:]\s*(<u>[^<]*</u>)',
            rf'\1：{emergency_name}', result
        )
    if emergency_phone:
        result = re.sub(
            r'(紧急联系?人\s*(?:电话|手机)\s*[：:])\s*(<u>[^<]*</u>|\s+)',
            rf'\1{emergency_phone}', result
        )
    if emergency_address:
        result = re.sub(
            r'(紧急联系?人\s*地址\s*[：:])\s*(<u>[^<]*</u>|\s+)',
            rf'\1{emergency_address}', result
        )

    result = re.sub(r'<u>\s*</u>', '', result)

    if signature_data:
        sig = f'<br><img src="{signature_data}" style="max-width:200px;max-height:100px;" alt="签名">'
        result = re.sub(
            r'(乙方[\(（]?.*?[\)）]?\s*签字)\s*[：:]\s*',
            lambda m: m.group(1) + '：' + sig, result,
        )
        if signature_data.split('/')[-1].split(',')[0] not in result:
            for pat, rep in [
                ('乙方（承揽人）签字：', f'乙方（承揽人）签字：{sig}'),
                ('乙方(签字):', f'乙方(签字):{sig}'),
                ('乙方(签字) :', f'乙方(签字) ：{sig}'),
                ('乙方（签字）', f'乙方（签字）{sig}'),
            ]:
                if pat in result:
                    result = result.replace(pat, rep)
                    break

    return result
